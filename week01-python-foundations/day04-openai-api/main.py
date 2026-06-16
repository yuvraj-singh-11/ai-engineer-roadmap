from fastapi import FastAPI, UploadFile, HTTPException
from pathlib import Path
# pyrefly: ignore [missing-import]
from pypdf import PdfReader
import docx
import os

app = FastAPI()

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)


@app.post("/upload")
async def upload_document(file: UploadFile):

    # Prevent path traversal attacks
    safe_filename = Path(file.filename).name

    file_path = os.path.join(UPLOAD_DIR, safe_filename)

    with open(file_path, "wb") as f:
        f.write(await file.read())

    text = ""

    try:
        extension = Path(safe_filename).suffix.lower()

        if extension == ".pdf":
            reader = PdfReader(file_path)

            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted

            page_count = len(reader.pages)

        elif extension == ".docx":
            doc = docx.Document(file_path)

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            page_count = None

        else:
            raise HTTPException(
                status_code=400,
                detail="Only PDF and DOCX files are supported"
            )

        return {
            "message": "File uploaded successfully",
            "filename": safe_filename,
            "pages": page_count,
            "characters": len(text),
            "words": len(text.split()),
            "preview": text[:1000]
        }

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error reading file: {str(e)}"
        )


@app.get("/document-metadata/{filename}")
def get_metadata(filename: str):

    safe_filename = Path(filename).name
    file_path = os.path.join(UPLOAD_DIR, safe_filename)

    if not os.path.exists(file_path):
        raise HTTPException(
            status_code=404,
            detail="File not found"
        )

    extension = Path(safe_filename).suffix.lower()

    try:

        metadata = {
            "filename": safe_filename,
            "file_size_kb": round(
                os.path.getsize(file_path) / 1024,
                2
            ),
            "file_type": extension
        }

        if extension == ".pdf":
            reader = PdfReader(file_path)
            metadata["pages"] = len(reader.pages)

        return metadata

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error reading metadata: {str(e)}"
        )


@app.get("/document-content/{filename}")
def get_document_content(filename: str):

    safe_filename = Path(filename).name
    file_path = os.path.join(UPLOAD_DIR, safe_filename)

    if not os.path.exists(file_path):
        raise HTTPException(
            status_code=404,
            detail="File not found"
        )

    extension = Path(safe_filename).suffix.lower()

    text = ""

    try:

        if extension == ".pdf":
            reader = PdfReader(file_path)

            for page in reader.pages:
                extracted = page.extract_text()

                if extracted:
                    text += extracted

        elif extension == ".docx":
            doc = docx.Document(file_path)

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

        else:
            raise HTTPException(
                status_code=400,
                detail="Unsupported file type"
            )

        return {
            "filename": safe_filename,
            "content": text
        }

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error reading document: {str(e)}"
        )
@app.get("/document-stats/{filename}")
def get_document_stats(filename: str):

    safe_filename = Path(filename).name
    file_path = os.path.join(UPLOAD_DIR, safe_filename)
    
    if not os.path.exists(file_path):
        raise HTTPException(
            status_code=404,
            detail="File not found"
        )
    
    extension = Path(safe_filename).suffix.lower()

    try:

        metadata = {
            "filename": safe_filename,
            "file_size_kb": round(
                os.path.getsize(file_path) / 1024,
                2
            ),
            "file_type": extension
        }

        if extension == ".pdf":
            reader = PdfReader(file_path)
            metadata["pages"] = len(reader.pages)

        return metadata

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error reading metadata: {str(e)}"
        )   